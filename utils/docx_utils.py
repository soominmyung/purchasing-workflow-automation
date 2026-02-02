"""
Save markdown/plain text to Word (.docx).
Output folders: output/analysis, output/pr, output/email_draft; or output/temp when use_temp_for_output.
Markdown pipe tables (| a | b |) are converted to real Word tables.
markdown_to_docx_bytes() builds .docx in memory (no disk) for stream/base64 download.
"""
import io
import re
import time
from pathlib import Path

from docx import Document

from config import settings

OUTPUT_ROOT = Path(__file__).resolve().parent.parent / "output"
ANALYSIS_DIR = OUTPUT_ROOT / "analysis"
PR_DIR = OUTPUT_ROOT / "pr"
EMAIL_DRAFT_DIR = OUTPUT_ROOT / "email_draft"
TEMP_DIR = OUTPUT_ROOT / "temp"


def _ensure_dirs() -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    PR_DIR.mkdir(parents=True, exist_ok=True)
    EMAIL_DRAFT_DIR.mkdir(parents=True, exist_ok=True)
    TEMP_DIR.mkdir(parents=True, exist_ok=True)


def cleanup_temp_output() -> None:
    """Delete files in output/temp older than temp_output_max_age_minutes (HF Space: avoid accumulation)."""
    if not TEMP_DIR.is_dir():
        return
    now = time.time()
    max_age_sec = getattr(settings, "temp_output_max_age_minutes", 5) * 60
    for f in TEMP_DIR.iterdir():
        if f.is_file() and (now - f.stat().st_mtime) > max_age_sec:
            try:
                f.unlink()
            except OSError:
                pass


def _sanitize_filename(name: str) -> str:
    """Remove characters invalid in filenames."""
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", "_", name).strip("._")
    return name or "unknown"


def _is_table_row(line: str) -> bool:
    """True if line looks like a markdown table row: | cell | cell |."""
    s = line.strip()
    return len(s) >= 2 and s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _is_separator_row(line: str) -> bool:
    """True if line is markdown table separator: | --- | --- |."""
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    inner = s[1:-1].strip()
    # Only pipes, dashes, spaces, colons
    return bool(re.match(r"^[\s|\-:]+$", inner))


def _parse_table_row(line: str) -> list[str]:
    """Split a table row into cell strings (strip each)."""
    s = line.strip()
    if not s.startswith("|") or not s.endswith("|"):
        return []
    parts = s[1:-1].split("|")
    return [p.strip() for p in parts]


def _add_table_to_doc(doc: Document, rows: list[list[str]]) -> None:
    """Add a Word table with the given rows. First row is header (bold)."""
    if not rows:
        return
    num_rows = len(rows)
    num_cols = max(len(r) for r in rows) if rows else 0
    if num_cols == 0:
        return
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    # Bold header row
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
    doc.add_paragraph()


def _content_to_docx_paragraphs(doc: Document, content: str) -> None:
    """Convert markdown/plain to docx: headings, paragraphs, and real Word tables for pipe tables."""
    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                row_line = lines[i]
                if _is_separator_row(row_line):
                    i += 1
                    continue
                table_rows.append(_parse_table_row(row_line))
                i += 1
            if table_rows:
                _add_table_to_doc(doc, table_rows)
            continue
        doc.add_paragraph(stripped)
        i += 1


def save_markdown_to_docx(file_path: Path, content: str) -> Path:
    """마크다운 내용을 .docx로 저장."""
    doc = Document()
    _content_to_docx_paragraphs(doc, content)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(file_path))
    return file_path


def markdown_to_docx_bytes(content: str) -> bytes:
    """Build .docx in memory (no disk). For stream/base64 download (HF Space, Framer)."""
    doc = Document()
    _content_to_docx_paragraphs(doc, content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_analysis_docx(snapshot_date: str, supplier: str, markdown_content: str) -> str:
    """
    output/analysis/ or output/temp (when use_temp_for_output)에 저장.
    반환: 저장된 파일 경로(문자열).
    """
    _ensure_dirs()
    if getattr(settings, "use_temp_for_output", False):
        cleanup_temp_output()
    safe_supplier = _sanitize_filename(supplier)
    filename = f"analysis_{snapshot_date}_{safe_supplier}.docx"
    dir_ = TEMP_DIR if getattr(settings, "use_temp_for_output", False) else ANALYSIS_DIR
    path = dir_ / filename
    save_markdown_to_docx(path, markdown_content)
    return str(path)


def save_pr_docx(snapshot_date: str, supplier: str, markdown_content: str) -> str:
    """
    output/pr/ or output/temp (when use_temp_for_output)에 저장.
    반환: 저장된 파일 경로(문자열).
    """
    _ensure_dirs()
    if getattr(settings, "use_temp_for_output", False):
        cleanup_temp_output()
    safe_supplier = _sanitize_filename(supplier)
    filename = f"pr_{snapshot_date}_{safe_supplier}.docx"
    dir_ = TEMP_DIR if getattr(settings, "use_temp_for_output", False) else PR_DIR
    path = dir_ / filename
    save_markdown_to_docx(path, markdown_content)
    return str(path)


def save_email_draft_docx(snapshot_date: str, supplier: str, text_content: str) -> str:
    """
    output/email_draft/ or output/temp (when use_temp_for_output)에 저장.
    반환: 저장된 파일 경로(문자열).
    """
    _ensure_dirs()
    if getattr(settings, "use_temp_for_output", False):
        cleanup_temp_output()
    safe_supplier = _sanitize_filename(supplier)
    filename = f"email_draft_{snapshot_date}_{safe_supplier}.docx"
    dir_ = TEMP_DIR if getattr(settings, "use_temp_for_output", False) else EMAIL_DRAFT_DIR
    path = dir_ / filename
    save_markdown_to_docx(path, text_content)
    return str(path)
